"""
Generate Employee Kiosk System User Manual as DOCX and PDF.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = r"c:\laragon\www\claude-technical-writer\output"
DOCX_PATH = os.path.join(OUTPUT_DIR, "Employee_Kiosk_User_Manual.docx")
PDF_PATH  = os.path.join(OUTPUT_DIR, "Employee_Kiosk_User_Manual.pdf")


# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set background colour of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p  = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pb.append(bottom)
    pPr.append(pb)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run()
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x7A)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    p.add_run(text)
    return p


def add_warning(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Warning: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run(text)
    return p


def add_screenshot(doc, filename, caption, width_inches=5.5):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        doc.add_paragraph(f"[Screenshot not found: {filename}]").italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    # style the caption runs that haven't been formatted yet
    if not cap.runs:
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_simple_table(doc, headers, rows, header_bg="1A567A", col_widths=None):
    """Add a table with a styled header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)

    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()   # spacing after table
    return table


# ── page setup ────────────────────────────────────────────────────────────────

def set_page_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)


# ── content sections ──────────────────────────────────────────────────────────

def add_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Employee Kiosk\nHeadcount Management System")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x7A)

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run("Comprehensive User Manual with Screenshots")
    sub.font.size = Pt(16)
    sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Version 1.0  ·  Last Updated: 2026-04-29\n"
        "https://headcount.employeeskiosk.com"
    ).font.size = Pt(11)

    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        ("1.", "System Overview"),
        ("2.", "Getting Started — Login"),
        ("3.", "Dashboard Overview"),
        ("4.", "Site Management"),
        ("5.", "Company and Client Management"),
        ("6.", "Headcount Management"),
        ("7.", "Location Management"),
        ("8.", "Reports and Analytics"),
        ("   8.1", "Headcount Snapshots"),
        ("   8.2", "Client Headcount Report"),
        ("   8.3", "TV Map Display"),
        ("9.", "Activity Log"),
        ("10.", "Customization and Settings"),
        ("11.", "User Management"),
        ("12.", "Troubleshooting and FAQs"),
        ("13.", "Best Practices"),
        ("14.", "Glossary"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{num}  {title}")
        run.font.size = Pt(11)
        if not num.startswith(" "):
            run.bold = True

    doc.add_page_break()


def section_1(doc):
    add_heading(doc, "1. System Overview", 1)
    add_body(doc,
        "The Employee Kiosk Headcount Management System is a web-based application for "
        "tracking and managing personnel headcount across distributed locations. It enables "
        "organizations to:"
    )
    bullets = [
        "Track total personnel headcount across multiple sites in real time",
        "Manage organizational structure including companies, clients, and sites",
        "Monitor active and floating personnel assignments",
        "Generate detailed headcount reports and analytics",
        "Maintain a full audit trail of all system changes",
        "Display live headcount data on TV screens for operational monitoring",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    add_heading(doc, "Who Should Use This System", 2)
    add_simple_table(doc,
        ["Role", "Description"],
        [
            ["Administrator / Webmaster", "Full access — manages all data, users, and configurations"],
            ["Site Manager", "Oversees headcount and assignments for assigned sites"],
            ["Company Manager", "Views personnel metrics across sites for their company"],
            ["Executive Management", "High-level metrics and trend analysis"],
            ["Finance / HR Personnel", "Generates reports and analyzes headcount data"],
        ],
        col_widths=[2.2, 4.0],
    )

    add_heading(doc, "Key Features at a Glance", 2)
    features = [
        "Real-time analytics dashboard with key metrics",
        "Site management with geographic coordinates",
        "Company and client relationship management",
        "Flexible headcount adjustment (addition, subtraction, floating)",
        "Location database with 1,599+ Philippine cities and municipalities",
        "Historical snapshot and trend tracking",
        "Client-specific headcount reporting",
        "TV display mode for real-time monitoring",
        "Complete activity audit trail",
        "CSV data export",
        "20+ customizable UI themes",
    ]
    for f in features:
        p = doc.add_paragraph(f, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def section_2(doc):
    add_heading(doc, "2. Getting Started — Login", 1)

    add_heading(doc, "Step 1: Open Your Web Browser", 2)
    add_body(doc, "Open Chrome, Firefox, Safari, or Edge (version 90+).")

    add_heading(doc, "Step 2: Navigate to the Application", 2)
    add_body(doc, "Enter the following URL in your browser address bar:")
    url_p = doc.add_paragraph("https://headcount.employeeskiosk.com")
    for run in url_p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

    add_heading(doc, "Step 3: Log In", 2)
    add_body(doc,
        "The login page appears with a clean, minimalist interface. Enter your credentials "
        "and click LOGIN to authenticate."
    )
    add_screenshot(doc, "screenshot-1777423409828.png",
                   "Figure 1.1 — Login Page: Enter your email and password to access the system")

    add_heading(doc, "Login Page Elements", 3)
    add_simple_table(doc,
        ["Element", "Description"],
        [
            ['"Welcome back!" heading', "Greeting displayed at the top of the login page"],
            ["Email address field", "Enter the email address associated with your account"],
            ["Password field", "Enter your password (characters are hidden)"],
            ["LOGIN button", "Click to submit credentials and log in"],
            ["FORGOT PASSWORD? link", "Use if you cannot remember your password"],
        ],
        col_widths=[2.2, 4.0],
    )

    add_heading(doc, "Step 4: Authentication", 2)
    add_body(doc,
        "Upon successful authentication you will be redirected to the Analytics Dashboard. "
        "If login fails, verify your email and password, check that Caps Lock is off, or use "
        "the FORGOT PASSWORD? link."
    )

    add_heading(doc, "First-Time User Setup", 2)
    add_heading(doc, "Understanding Your Role", 3)
    add_simple_table(doc,
        ["Role", "Access Level"],
        [
            ["Webmaster / Administrator", "Full access to all system features"],
            ["Manager", "View and report on assigned areas; limited editing"],
            ["Employee", "Read-only access — view dashboards and reports only"],
        ],
        col_widths=[2.5, 3.7],
    )
    add_note(doc, "Contact your system administrator if you need a different access level.")

    add_heading(doc, "Changing Your Password", 3)
    steps = [
        "Click your profile icon (your initials) in the top-right corner",
        "Select Settings from the dropdown menu",
        "Follow the prompts to create a new password",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        p.paragraph_format.space_after = Pt(2)

    add_note(doc, "Change your default password on first login for security.")
    doc.add_page_break()


def section_3(doc):
    add_heading(doc, "3. Dashboard Overview", 1)
    add_body(doc,
        "The Analytics Dashboard is the first screen you see after logging in. It provides "
        "a real-time summary of headcount across all sites and companies."
    )

    add_heading(doc, "Dashboard — Main View", 2)
    add_screenshot(doc, "screenshot-1777423540101.png",
                   "Figure 2.1 — Analytics Dashboard Overview: Real-time headcount metrics and charts")

    add_heading(doc, "Key Elements", 3)
    add_simple_table(doc,
        ["Element", "Description"],
        [
            ["Total Headcount card", "Total personnel across all sites (e.g., 403)"],
            ["Total Sites card", "Number of operational locations (e.g., 12)"],
            ["Regional Distribution", "Headcount and percentage for Luzon, Visayas, Mindanao"],
            ["Headcount by Region chart", "Dual-axis: headcount (left axis) vs. site count (right axis)"],
            ["Headcount by Company chart", "Donut chart showing percentage share per company"],
            ["User profile area", "Name, role, quick access to Settings, Profile, Support, Logout"],
        ],
        col_widths=[2.5, 3.7],
    )

    add_heading(doc, "Dashboard — Full Scrolled View", 2)
    add_screenshot(doc, "screenshot-1777424950515.png",
                   "Figure 2.2 — Dashboard Full View: Headcount Trend filters and Sites Overview table")

    add_heading(doc, "Reading the Charts", 2)
    add_heading(doc, "Headcount by Region Chart", 3)
    add_body(doc,
        "This dual-axis chart shows headcount on the left axis (line) and site count on the "
        "right axis (bars). The X-axis shows regions or time periods. Use it to identify "
        "which regions have the most personnel and sites."
    )

    add_heading(doc, "Headcount by Company Donut Chart", 3)
    add_body(doc,
        "Each segment represents one company. Segment size reflects relative headcount. "
        "Hover over a segment to see exact numbers. Example: Prime (167), Primus (211), Niupro (25)."
    )

    add_heading(doc, "Using Time-Period Filters", 2)
    add_body(doc,
        "Click 7 Days, 30 Days, 90 Days, or All above the Headcount Trend chart. The chart "
        "updates to show data for that period. Use this to identify recent changes or long-term trends."
    )

    add_heading(doc, "Sites Overview Table", 2)
    add_simple_table(doc,
        ["Column", "Description"],
        [
            ["Site Name", "Name of the location"],
            ["Client", "Client organization served at that site"],
            ["Company", "Company managing the site"],
            ["Region", "Geographic region (Luzon, Visayas, Mindanao)"],
            ["Headcount", "Total personnel at the site"],
        ],
        col_widths=[2.0, 4.2],
    )
    add_body(doc,
        "Use the search box to filter by site name, the region dropdown to filter by area, "
        "and click any column header to sort. Use page controls for large datasets."
    )
    doc.add_page_break()


def section_4(doc):
    add_heading(doc, "4. Site Management", 1)
    add_body(doc,
        "Site Management is where you view, create, edit, and delete all operational site "
        "locations. Each site represents a physical location where personnel are assigned."
    )

    add_heading(doc, "Accessing Site Management", 2)
    steps = [
        "In the sidebar, click HEADCOUNT to expand the submenu",
        "Click Site Management",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Site Management Interface", 2)
    add_screenshot(doc, "screenshot-1777425550788.png",
                   "Figure 3.1 — Site Management Page: Full sites table with headcount, "
                   "geographic coordinates, and action buttons")

    add_heading(doc, "Summary Metrics", 3)
    add_body(doc,
        "At the top of the page: Total Headcount 403 (Active: 400, Floating: 3). "
        "Company distribution: Prime 167 / 3 sites · Primus 211 / 8 sites · Niupro 25 / 1 site."
    )

    add_heading(doc, "Sites Table Columns", 2)
    add_simple_table(doc,
        ["Column", "Description"],
        [
            ["Site Name", "Unique name of the location"],
            ["Client", "Client organization served at this site"],
            ["Company", "Company managing the site"],
            ["Region", "Geographic region"],
            ["Headcount", "Total personnel (Active + Floating)"],
            ["Active", "Personnel actively assigned to the site"],
            ["Floating", "Personnel temporarily unassigned"],
            ["Latitude", "North-South geographic coordinate"],
            ["Longitude", "East-West geographic coordinate"],
            ["Actions", "Edit and Delete buttons"],
        ],
        col_widths=[1.8, 4.4],
    )

    add_heading(doc, "Adding a New Site", 2)
    steps = [
        "Click New Site (top-right of the page)",
        "Enter Site Name, Client, Company, and Region (all required)",
        "Enter Latitude and Longitude (find with Google Maps)",
        "Click Save — the site appears in the table immediately",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Editing a Site", 2)
    steps = [
        "Find the site in the table",
        "Click Edit in the Actions column",
        "Modify the desired fields",
        "Click Save",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Deleting a Site", 2)
    steps = [
        "Find the site in the table",
        "Click Delete in the Actions column",
        "Confirm the deletion in the dialog",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)
    add_warning(doc,
        "Deleting a site is permanent and removes all associated data. "
        "The action is recorded in the Activity Log."
    )

    add_heading(doc, "Understanding Geographic Coordinates", 2)
    add_simple_table(doc,
        ["Coordinate", "Measures", "Positive =", "Negative ="],
        [
            ["Latitude", "North-South position", "North", "South"],
            ["Longitude", "East-West position", "East", "West"],
        ],
        col_widths=[1.5, 2.0, 1.5, 1.5],
    )
    add_body(doc,
        "Example: Latitude 14.5564 + Longitude 121.0177 = a location in Metro Manila."
    )
    doc.add_page_break()


def section_5(doc):
    add_heading(doc, "5. Company and Client Management", 1)

    add_heading(doc, "Organizational Hierarchy", 2)
    add_body(doc, "The system uses a three-level hierarchy:")
    hierarchy = [
        "Companies (top level) — operating organizations, e.g., Prime, Primus, Niupro",
        "Clients (middle level) — organizations served by companies, e.g., BDO Unibank, HUAWEI",
        "Sites (operational level) — physical locations where services are delivered",
    ]
    for h in hierarchy:
        p = doc.add_paragraph(h, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    add_heading(doc, "Accessing Company and Client Management", 2)
    steps = [
        "In the sidebar, click HEADCOUNT",
        "Click Company & Clients",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Company and Client Management Interface", 2)
    add_screenshot(doc, "screenshot-1777426048222.png",
                   "Figure 4.1 — Company and Client Management: Three companies, "
                   "fifteen clients, with tab navigation")

    add_heading(doc, "Summary Metrics", 3)
    add_body(doc,
        "Top of page: 3 Companies · 15 Clients · 5.0 Average Clients per Company. "
        "Use the Companies and Clients tabs to switch between the two views."
    )

    add_heading(doc, "Current Companies", 3)
    add_simple_table(doc,
        ["Company", "Clients", "Created"],
        [
            ["Niupro", "5", "3/30/2026"],
            ["Prime", "5", "3/30/2026"],
            ["Primus", "5", "3/30/2026"],
        ],
        col_widths=[2.5, 1.5, 2.2],
    )

    add_heading(doc, "Adding a New Company", 2)
    steps = [
        "Click New Company",
        "Enter the Company Name (required) and optional Description",
        "Click Save",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Adding a New Client", 2)
    steps = [
        "Click the Clients tab",
        "Click New Client",
        "Enter the Client Name and select the managing Company",
        "Click Save",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_note(doc,
        "You cannot delete a company that still has associated clients or sites. "
        "Remove those dependencies first."
    )
    doc.add_page_break()


def section_6(doc):
    add_heading(doc, "6. Headcount Management", 1)

    add_heading(doc, "Headcount Types", 2)
    add_simple_table(doc,
        ["Type", "Description"],
        [
            ["Active", "Personnel currently assigned to a site with full accountability"],
            ["Floating", "Personnel temporarily unassigned or in transition between sites"],
        ],
        col_widths=[1.8, 4.4],
    )
    add_body(doc,
        "Both types count toward total headcount but are tracked separately for planning purposes."
    )

    add_heading(doc, "Accessing Headcount Adjustments", 2)
    steps = [
        "Click HEADCOUNT in the sidebar",
        "Click Headcount Adjustments",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Headcount Adjustments Interface", 2)
    add_screenshot(doc, "screenshot-1777426118507.png",
                   "Figure 5.1 — Headcount Adjustments Page: Adjustment types and full "
                   "adjustment history with date, type, company, and site details")

    add_heading(doc, "Adjustment Types", 2)
    add_simple_table(doc,
        ["Type", "Effect on Site Headcount", "When to Use"],
        [
            ["Addition", "Increases total and active count", "New hire, incoming transfer"],
            ["Subtraction", "Decreases total and active count", "Resignation, outgoing transfer"],
            ["Floating", "Active count decreases; floating count increases; total unchanged",
             "Personnel in transition, on leave, or on standby"],
        ],
        col_widths=[1.4, 2.8, 2.0],
    )

    add_heading(doc, "Making a Headcount Adjustment", 2)
    steps = [
        "Click Add Headcount",
        "Select Type (Addition, Subtraction, or Floating)",
        "Select Company, Client, and Site",
        "Enter the Amount (number of personnel affected)",
        "Set the Date and add an optional Note explaining the reason",
        "Click Submit",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_body(doc,
        "The adjustment appears immediately in the Adjustment History table and updates "
        "all dashboard metrics, the Activity Log, and trend snapshots."
    )

    add_heading(doc, "Adjustment History Table", 2)
    add_simple_table(doc,
        ["Column", "Description"],
        [
            ["Date", "When the adjustment was recorded"],
            ["Type", "Addition, Subtraction, or Floating"],
            ["Company", "Company affected"],
            ["Client", "Client affected"],
            ["Site", "Site affected"],
            ["Addition", "Number added (if applicable)"],
            ["Subtraction", "Number subtracted (if applicable)"],
        ],
        col_widths=[1.8, 4.4],
    )
    doc.add_page_break()


def section_7(doc):
    add_heading(doc, "7. Location Management", 1)
    add_body(doc,
        "The system contains a complete Philippines location database used for site creation, "
        "address validation, and geographic coordinate reference."
    )

    add_heading(doc, "Accessing Location Management", 2)
    steps = [
        "Click HEADCOUNT in the sidebar",
        "Click Location Management",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Location Management Interface", 2)
    add_screenshot(doc, "screenshot-1777426291369.png",
                   "Figure 6.1 — Location Management Page: 1,599 Philippine cities and "
                   "municipalities with coordinates and bulk-delete capability")

    add_heading(doc, "Database Summary", 3)
    add_body(doc,
        "1,599 total cities and municipalities · 141 distinct provinces. "
        "Covers the entire Philippines."
    )

    add_heading(doc, "Location Table Fields", 2)
    add_simple_table(doc,
        ["Field", "Description", "Example"],
        [
            ["Province/Municipality", "Regional division", "Abra"],
            ["City/Municipality", "Specific city or municipality", "Bangued"],
            ["Latitude", "North-South coordinate", "17.5965"],
            ["Longitude", "East-West coordinate", "120.6179"],
        ],
        col_widths=[2.0, 2.5, 1.7],
    )

    add_heading(doc, "Searching for Locations", 2)
    add_body(doc,
        "Type the city, municipality, or province name in the search field. "
        "The table filters in real time. Examples: searching 'Manila' shows all "
        "Manila-related locations; searching 'Quezon' shows Quezon City and Quezon Province."
    )

    add_heading(doc, "Adding a New Location", 2)
    steps = [
        "Click New Location",
        "Enter Province/Municipality, City/Municipality, Latitude, and Longitude",
        "Click Save",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)
    add_note(doc, "Coordinates can be found by right-clicking a location in Google Maps.")

    add_heading(doc, "Bulk Delete", 2)
    steps = [
        "Check the checkboxes beside the locations to remove",
        "Click Delete Selected",
        "Confirm the deletion",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)
    add_warning(doc, "Only delete locations that are not currently assigned to any site.")
    doc.add_page_break()


def section_8(doc):
    add_heading(doc, "8. Reports and Analytics", 1)
    add_body(doc,
        "The system provides three dedicated report and display tools in addition to the "
        "main Analytics Dashboard: Headcount Snapshots, Client Headcount Report, and TV Map Display."
    )

    # 8.1
    add_heading(doc, "8.1 Headcount Snapshots", 2)
    add_body(doc,
        "A snapshot captures the exact headcount state at a specific point in time. "
        "Snapshots are used for trend analysis, historical comparison, and date-specific reporting."
    )

    add_heading(doc, "Accessing Headcount Snapshots", 3)
    steps = [
        "Click HEADCOUNT in the sidebar",
        "Click Headcount Snapshots",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Headcount Snapshots Interface", 3)
    add_screenshot(doc, "screenshot-1777426474070.png",
                   "Figure 7.1 — Headcount Snapshots Page: Grand total, company breakdown, "
                   "trend filters, and daily snapshot history table")

    add_heading(doc, "Page Elements", 3)
    add_simple_table(doc,
        ["Element", "Description"],
        [
            ["Grand Total (Latest)", "Current total headcount (e.g., 403 as of 2026-04-28)"],
            ["Company breakdown", "Prime: 167 · Primus: 211 · Niupro: 25"],
            ["Headcount Trend", "Chart with Apply/Clear date-range filters"],
            ["Take Snapshot Now", "Button to capture current headcount immediately"],
            ["Export CSV", "Download full snapshot history as a spreadsheet"],
            ["Daily Snapshots table", "Columns: Date · Prime · Primus · Niupro"],
        ],
        col_widths=[2.2, 4.0],
    )

    add_heading(doc, "Taking a Manual Snapshot", 3)
    add_body(doc,
        "Click Take Snapshot Now. The system records the current headcount with a timestamp "
        "and the new entry appears at the top of the Daily Snapshots table. "
        "The system also takes automatic daily snapshots."
    )

    add_heading(doc, "Filtering Snapshots by Date", 3)
    steps = [
        "Enter a start and end date in the Headcount Trend date filters",
        "Click Apply — the table and chart update",
        "Click Clear to reset and view all snapshots",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # 8.2
    add_heading(doc, "8.2 Client Headcount Report", 2)
    add_body(doc,
        "Generate a report showing headcount by client as of any specific past or present date — "
        "useful for historical lookups and client billing."
    )

    add_heading(doc, "Accessing Client Headcount Report", 3)
    steps = [
        "Click HEADCOUNT in the sidebar",
        "Click Client HC Report",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Client Headcount Report Interface", 3)
    add_screenshot(doc, "screenshot-1777426520745.png",
                   "Figure 8.1 — Client Headcount Report Page: Date picker, company filter, "
                   "and results table showing headcount per client")

    add_heading(doc, "Generating a Report", 3)
    steps = [
        "Click the As-Of Date field and select the desired date",
        "Optionally select a company in Filter by Company to narrow results",
        "Click Generate Report",
        "The results table populates with client headcount for the selected date",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Report Columns", 3)
    add_simple_table(doc,
        ["Column", "Description"],
        [
            ["Client", "Client organization name"],
            ["Company", "Company managing the client"],
            ["Headcount", "Total personnel assigned to this client"],
            ["Active", "Actively assigned personnel (excludes floating)"],
        ],
        col_widths=[1.8, 4.4],
    )

    doc.add_paragraph()

    # 8.3
    add_heading(doc, "8.3 TV Map Display", 2)
    add_body(doc,
        "Display live headcount information on a large monitor or TV for control room "
        "monitoring, lobby displays, or executive dashboards."
    )

    add_heading(doc, "Accessing TV Map Display", 3)
    steps = [
        "Click HEADCOUNT in the sidebar",
        "Click TV Map Display",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "TV Map Display Interface", 3)
    add_screenshot(doc, "screenshot-1777426557255.png",
                   "Figure 9.1 — TV Map Display: Real-time headcount tracker with "
                   "LIVE indicator, total manpower, company cards, and site listings")

    add_heading(doc, "Display Elements", 3)
    add_simple_table(doc,
        ["Element", "Description"],
        [
            ["LIVE indicator", "Green badge confirming real-time data updates"],
            ["TOTAL MANPOWER", "Grand total headcount across all sites (e.g., 403)"],
            ["Current time and date", "Live clock and full date"],
            ["Last updated timestamp", "When the data was last refreshed"],
            ["Company cards", "Headcount and site count per company"],
            ["Site listing", "Site name · Client · Personnel count"],
            ["+/− navigation buttons", "Browse through all sites in the system"],
        ],
        col_widths=[2.2, 4.0],
    )

    add_heading(doc, "Setting Up TV Display", 3)
    steps = [
        "Connect a computer to the TV or large monitor",
        "Open a browser and navigate to TV Map Display",
        "The display auto-updates approximately every 1-2 minutes",
        "Use the +/- buttons to navigate through site listings",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def section_9(doc):
    add_heading(doc, "9. Activity Log", 1)
    add_body(doc,
        "The Activity Log records every change made in the system — who did it, what was "
        "changed, and when. It provides a complete audit trail for compliance, accountability, "
        "and troubleshooting."
    )

    add_heading(doc, "Accessing the Activity Log", 2)
    steps = [
        "Click HEADCOUNT in the sidebar",
        "Click Activity Log",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Activity Log Interface", 2)
    add_screenshot(doc, "screenshot-1777426451609.png",
                   "Figure 10.1 — Activity Log Page: 28 records with action type filters "
                   "and full details including IP address and timestamp")

    add_heading(doc, "Summary Metrics", 3)
    add_simple_table(doc,
        ["Metric", "Value", "Description"],
        [
            ["Total Records", "28", "All events recorded in the system"],
            ["Created", "12", "New records added to the system"],
            ["Updated", "16", "Existing records modified"],
            ["Deleted", "0", "Records removed from the system"],
        ],
        col_widths=[2.0, 1.5, 2.7],
    )

    add_heading(doc, "Activity Log Table Columns", 2)
    add_simple_table(doc,
        ["Column", "Description"],
        [
            ["Action", "What was done (Create, Update, Delete, etc.)"],
            ["Type", "Category of change (Site, Company, Headcount, etc.)"],
            ["Description", "Specific detail of what changed"],
            ["Performed By", "User who made the change"],
            ["IP Address", "Network address of the user at time of change"],
            ["Timestamp", "Date and time the action was performed"],
        ],
        col_widths=[1.8, 4.4],
    )

    add_heading(doc, "Filtering the Activity Log", 2)
    add_body(doc,
        "Use the All Actions dropdown to filter by action type (Create, Update, Delete). "
        "Use the All Types dropdown to filter by category (Site, Company, Headcount, Location, User)."
    )

    add_heading(doc, "Example Log Entries", 2)
    add_simple_table(doc,
        ["Action", "Type", "Description", "Performed By"],
        [
            ["Create", "Site", "New site 'BGC' created", "Amy Elsner"],
            ["Update", "Headcount", "Added 5 to Prime at BGC", "Amy Elsner"],
            ["Update", "Site", "Edited coordinates for bdo luzon", "Amy Elsner"],
            ["Delete", "Client", "Removed test client", "Admin User"],
        ],
        col_widths=[1.2, 1.3, 2.8, 1.5],
    )
    doc.add_page_break()


def section_10(doc):
    add_heading(doc, "10. Customization and Settings", 1)

    add_heading(doc, "Accessing Settings", 2)
    add_body(doc,
        "Method 1: Click the Settings gear icon (⚙) in the top navigation bar. "
        "Method 2: Click your profile icon (top-right) → Settings."
    )

    add_heading(doc, "Settings Interface", 2)
    add_screenshot(doc, "screenshot-1777423801316.png",
                   "Figure 11.1 — Settings Panel: Theme thumbnails, Light/Dark mode toggle, "
                   "and sidebar layout options")

    add_heading(doc, "Color Themes", 2)
    add_body(doc, "The system offers 25+ themes organized by color family:")
    add_simple_table(doc,
        ["Category", "Available Themes"],
        [
            ["Neutral", "Noir, Slate, Gray, Zinc, Neutral, Stone"],
            ["Warm", "Amber, Orange, Yellow, Rose, Pink, Fuchsia"],
            ["Cool", "Teal, Cyan, Sky, Blue, Indigo, Violet, Purple"],
            ["Green", "Emerald, Green, Lime"],
            ["Special", "Soho, Viva, Ocean"],
        ],
        col_widths=[1.8, 4.4],
    )
    add_body(doc,
        "To change the theme: open Settings, click any theme thumbnail, and the interface "
        "updates immediately — no save button needed."
    )

    add_heading(doc, "Light / Dark Mode", 2)
    add_body(doc,
        "Toggle between Light Mode (light background, dark text) and Dark Mode "
        "(dark background, light text). The setting persists across sessions."
    )

    add_heading(doc, "Sidebar Layout Options", 2)
    add_simple_table(doc,
        ["Layout", "Description"],
        [
            ["Static", "Sidebar always visible on screen"],
            ["Overlay", "Sidebar slides over content when opened"],
            ["Slim", "Narrow sidebar showing icons only"],
            ["Slim+", "Narrow sidebar with expanded tooltips"],
            ["Reveal", "Hidden sidebar that slides out on hover"],
            ["Drawer", "Mobile-style drawer menu"],
        ],
        col_widths=[1.8, 4.4],
    )

    add_heading(doc, "Sidebar Position", 2)
    add_body(doc,
        "Choose Start (left side, default) or End (right side) for the sidebar position."
    )
    doc.add_page_break()


def section_11(doc):
    add_heading(doc, "11. User Management", 1)

    add_heading(doc, "User Roles and Permissions", 2)
    add_simple_table(doc,
        ["Role", "Access"],
        [
            ["Webmaster / Administrator",
             "Full access — create, edit, delete all data and manage users"],
            ["Manager",
             "View assigned areas and generate reports; limited editing capability"],
            ["Employee",
             "Read-only — view dashboards and reports only; cannot make data changes"],
        ],
        col_widths=[2.2, 4.0],
    )

    add_heading(doc, "Accessing User Management", 2)
    add_body(doc,
        "In the sidebar, click USER MANAGEMENT to expand: "
        "List (view all users) and Create (add a new user)."
    )

    add_heading(doc, "User Profile Dropdown", 2)
    add_screenshot(doc, "screenshot-1777423980601.png",
                   "Figure 11.2 — User Profile Dropdown: Quick access to Settings, "
                   "Profile, Support, and Logout")

    add_heading(doc, "Creating a New User", 2)
    steps = [
        "Navigate to USER MANAGEMENT > Create",
        "Enter First Name, Last Name, Email (must be unique), Password, and Confirm Password",
        "Select a Role: Webmaster, Manager, or Employee",
        "Click Create User — the account is created and a login email is sent",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Editing a User", 2)
    steps = [
        "Go to USER MANAGEMENT > List",
        "Find the user and click Edit",
        "Modify name, email, role, or status",
        "Click Save",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Deactivating a User", 2)
    steps = [
        "Find the user in the list",
        "Click Deactivate or toggle Status to Inactive",
        "Confirm the action",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)
    add_note(doc, "The user cannot log in after deactivation. No data is deleted.")

    add_heading(doc, "Resetting a User Password", 2)
    steps = [
        "Find the user in the list",
        "Click Reset Password",
        "The system sends a password reset email to the user",
        "The user follows the link to set a new password",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def section_12(doc):
    add_heading(doc, "12. Troubleshooting and FAQs", 1)

    add_heading(doc, "Cannot Login", 2)
    add_simple_table(doc,
        ["Possible Cause", "Solution"],
        [
            ["Incorrect email or password", "Retype carefully; verify Caps Lock is off"],
            ["Account not yet created", "Contact your system administrator"],
            ["Account deactivated", "Contact your system administrator"],
            ["Forgotten password",
             "Click FORGOT PASSWORD? on the login page and follow the email link"],
        ],
        col_widths=[2.5, 3.7],
    )

    add_heading(doc, "Blank Dashboard or Missing Data", 2)
    steps = [
        "Wait 5-10 seconds for the page to fully load",
        "Check your internet connection",
        "Refresh the page (F5 or Ctrl+R)",
        "Clear browser cache: Chrome → Settings → Privacy → Clear browsing data",
        "Try a different browser",
        "Contact your administrator if the issue persists",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Tables Not Displaying Correctly", 2)
    steps = [
        "Clear all filters (click Clear All or reset each filter)",
        "Verify you are in the correct section of the system",
        "Refresh the page",
        "Use Chrome, Firefox, Safari, or Edge version 90+",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Cannot Edit or Delete Records", 2)
    add_simple_table(doc,
        ["Possible Cause", "Solution"],
        [
            ["Insufficient role permissions",
             "Ask your administrator to update your role"],
            ["Record has dependent data",
             "Remove linked records first (e.g., remove clients before deleting a company)"],
            ["Browser JavaScript issue",
             "Try a different browser or clear the browser cache"],
        ],
        col_widths=[2.5, 3.7],
    )

    add_heading(doc, '"NOT FOUND" Error Page', 2)
    add_screenshot(doc, "screenshot-1777425276677.png",
                   "Figure 12.1 — 404 Not Found Error Page: "
                   "Click GO BACK TO DASHBOARD to return to the main interface")
    add_body(doc,
        "This page appears when you navigate to a URL that does not exist or a record that "
        "was deleted. Click GO BACK TO DASHBOARD to return to the main interface."
    )

    add_heading(doc, "Frequently Asked Questions", 2)
    faqs = [
        ("How often are headcount snapshots taken?",
         "Daily automatically. You can also take one manually via Headcount Snapshots > Take Snapshot Now."),
        ("Can I undo a headcount adjustment?",
         "There is no direct undo. Create an opposite adjustment — subtract what you added, or add back what you subtracted."),
        ("What is the difference between Active and Floating headcount?",
         "Active = personnel currently assigned to a site. Floating = personnel in transition or on standby. Both count toward the total."),
        ("Can I view headcount for a past date?",
         "Yes. Go to Client HC Report, set the As-Of Date to the desired historical date, and click Generate Report."),
        ("Who made a specific change in the system?",
         "Check the Activity Log — it shows the action, change type, description, and the user who performed it."),
        ("Can I access the system on a mobile device?",
         "Yes. The system is mobile-responsive and adapts to tablets and smartphones."),
        ("How do I change my password?",
         "Click your profile icon (top-right) > Settings > Change Password. If you forgot your password, use FORGOT PASSWORD? on the login page."),
        ("How do I print a report?",
         "Use your browser's print function (Ctrl+P / Cmd+P). Select 'Save as PDF' to create a digital copy."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(f"Q: {q}")
        run.bold = True
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        ans = doc.add_paragraph(f"A: {a}")
        ans.paragraph_format.left_indent = Cm(0.6)
        ans.paragraph_format.space_after = Pt(8)
        for r in ans.runs:
            r.font.size = Pt(10)

    doc.add_page_break()


def section_13(doc):
    add_heading(doc, "13. Best Practices", 1)

    add_heading(doc, "Data Management", 2)
    add_simple_table(doc,
        ["Practice", "Why", "How"],
        [
            ["Regular headcount audits",
             "Ensures data accuracy",
             "Review metrics weekly; compare Active vs. Floating; investigate discrepancies"],
            ["Consistent site naming",
             "Ensures data quality",
             "Use a standard naming convention; include company and region where possible"],
            ["Timely adjustments",
             "Keeps data current",
             "Record adjustments same day; add clear notes explaining the reason"],
            ["Regular snapshots",
             "Enables trend analysis",
             "Take manual snapshots weekly or monthly; export for external use"],
        ],
        col_widths=[1.8, 1.8, 2.7],
    )

    add_heading(doc, "User Management", 2)
    add_simple_table(doc,
        ["Practice", "Why", "How"],
        [
            ["Minimum necessary permissions",
             "Security",
             "Assign Webmaster only when full access is truly required"],
            ["Deactivate departing users",
             "Security",
             "Immediately deactivate accounts when personnel leave the organization"],
            ["Strong passwords",
             "Security",
             "Mix letters, numbers, and symbols; change on first login"],
            ["Regular user list review",
             "Accuracy",
             "Periodically audit users and deactivate stale accounts"],
        ],
        col_widths=[1.8, 1.5, 3.0],
    )

    add_heading(doc, "Reporting Best Practices", 2)
    bullets = [
        "Review the Analytics Dashboard daily or weekly to catch unusual changes early",
        "Generate Client Headcount Reports on a consistent schedule for accurate billing",
        "Export and archive reports monthly for record retention",
        "Review the Activity Log regularly to verify data integrity and track accountability",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    add_heading(doc, "Recommended Workflows", 2)
    add_heading(doc, "Adding a New Site", 3)
    steps = [
        "Confirm the physical location and determine company/client relationships",
        "Look up accurate coordinates from Google Maps (right-click → copy coordinates)",
        "Enter the site in the system and verify it appears on the dashboard",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    add_heading(doc, "Making a Headcount Adjustment", 3)
    steps = [
        "Receive request with site, company, and personnel details",
        "Verify the correct adjustment type (Addition, Subtraction, or Floating)",
        "Record in the system with a clear explanatory note",
        "Confirm the updated totals in Site Management and the dashboard",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}").paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def section_14(doc):
    add_heading(doc, "14. Glossary", 1)
    terms = [
        ("Active Headcount",
         "Personnel currently assigned to a site with full accountability."),
        ("Activity Log",
         "System record of all changes including who made them, what was changed, and when."),
        ("Analytics Dashboard",
         "Main screen showing real-time headcount metrics, charts, and the Sites Overview table."),
        ("As-Of Date",
         "Specific date for which a report is generated, showing system state at that point in time."),
        ("Audit Trail",
         "Complete history of all system actions for compliance and accountability purposes."),
        ("Client",
         "Organization served by a company at operational sites. Multiple clients per company."),
        ("Company",
         "Operating organization that manages personnel and sites (e.g., Prime, Primus, Niupro)."),
        ("CSV",
         "Comma-Separated Values — file format for data export, compatible with Excel."),
        ("Floating Headcount",
         "Personnel temporarily unassigned or in transition between sites."),
        ("Grand Total",
         "Total headcount across the entire organization."),
        ("Headcount Adjustment",
         "Process of adding, removing, or reclassifying personnel at a site."),
        ("Headcount Snapshot",
         "Point-in-time capture of headcount data used for trend analysis and historical reporting."),
        ("Latitude",
         "North-South geographic coordinate (positive values = North of equator)."),
        ("Longitude",
         "East-West geographic coordinate (positive values = East of prime meridian)."),
        ("Role",
         "User permission level determining accessible features: Webmaster, Manager, or Employee."),
        ("Site",
         "Physical location where personnel are assigned and services are delivered."),
        ("Snapshot",
         "See Headcount Snapshot."),
        ("Total Headcount",
         "Complete count of all personnel (active + floating) across the organization."),
        ("Trend",
         "Pattern of headcount changes over time, indicating increase, decrease, or stability."),
        ("TV Map Display",
         "Full-screen real-time display of live headcount information for monitors and TV screens."),
        ("Webmaster",
         "Administrative user role with full access to all system features and settings."),
    ]
    for term, definition in terms:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{term}: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x7A)
        p.add_run(definition)


def add_appendix(doc):
    doc.add_page_break()
    add_heading(doc, "Appendix: Screenshot Reference Index", 1)
    add_simple_table(doc,
        ["Figure", "Description", "File"],
        [
            ["1.1", "Login Page", "screenshot-1777423409828.png"],
            ["2.1", "Analytics Dashboard Overview", "screenshot-1777423540101.png"],
            ["2.2", "Dashboard Full View (Trend + Sites Table)", "screenshot-1777424950515.png"],
            ["3.1", "Site Management Page", "screenshot-1777425550788.png"],
            ["4.1", "Company and Client Management", "screenshot-1777426048222.png"],
            ["5.1", "Headcount Adjustments Page", "screenshot-1777426118507.png"],
            ["6.1", "Location Management Page", "screenshot-1777426291369.png"],
            ["7.1", "Headcount Snapshots Page", "screenshot-1777426474070.png"],
            ["8.1", "Client Headcount Report Page", "screenshot-1777426520745.png"],
            ["9.1", "TV Map Display", "screenshot-1777426557255.png"],
            ["10.1", "Activity Log Page", "screenshot-1777426451609.png"],
            ["11.1", "Settings / Theme Panel", "screenshot-1777423801316.png"],
            ["11.2", "User Profile Dropdown Menu", "screenshot-1777423980601.png"],
            ["12.1", "404 Not Found Error Page", "screenshot-1777425276677.png"],
        ],
        col_widths=[0.7, 3.2, 2.8],
    )

    doc.add_paragraph()
    add_heading(doc, "System Requirements", 2)
    add_simple_table(doc,
        ["Requirement", "Specification"],
        [
            ["Supported Browsers", "Chrome 90+, Firefox 88+, Safari 14+, Edge 90+"],
            ["Internet Connection", "Stable; minimum 2 Mbps recommended"],
            ["Devices", "Desktop, laptop, tablet, or smartphone"],
            ["Minimum Screen Width", "320 px"],
        ],
        col_widths=[2.2, 4.0],
    )

    add_heading(doc, "Keyboard Shortcuts", 2)
    add_simple_table(doc,
        ["Action", "Windows", "Mac"],
        [
            ["Refresh page", "F5 or Ctrl+R", "Cmd+R"],
            ["Clear cache", "Ctrl+Shift+Del", "Cmd+Shift+Delete"],
            ["Print", "Ctrl+P", "Cmd+P"],
            ["Find on page", "Ctrl+F", "Cmd+F"],
            ["New tab", "Ctrl+T", "Cmd+T"],
        ],
        col_widths=[2.5, 1.8, 1.8],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph(
        "Employee Kiosk Headcount Management System — User Manual  |  "
        "Version 1.0  |  2026-04-29"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── main ──────────────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    set_page_margins(doc)

    # default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    add_cover(doc)
    add_toc(doc)
    section_1(doc)
    section_2(doc)
    section_3(doc)
    section_4(doc)
    section_5(doc)
    section_6(doc)
    section_7(doc)
    section_8(doc)
    section_9(doc)
    section_10(doc)
    section_11(doc)
    section_12(doc)
    section_13(doc)
    section_14(doc)
    add_appendix(doc)

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")
    return DOCX_PATH


def build_pdf(docx_path):
    try:
        from docx2pdf import convert
        convert(docx_path, PDF_PATH)
        print(f"PDF  saved: {PDF_PATH}")
    except Exception as e:
        print(f"PDF conversion failed: {e}")
        print("Ensure Microsoft Word is installed and docx2pdf can access it.")


if __name__ == "__main__":
    docx_path = build_docx()
    build_pdf(docx_path)
